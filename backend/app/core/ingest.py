"""Document ingestion: PDF / DOCX / image / text -> chunks -> embeddings -> Chroma.

Images (dropped screenshots) are transcribed by the vision model so their
content becomes searchable text — that's how the client "trains by screenshot".
"""
from __future__ import annotations

from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.config import settings
from app.core import vectorstore
from app.core.db import tx
from app.core.ollama_client import ollama

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}
TEXT_EXTS = {".txt", ".md"}

_OCR_PROMPT = (
    "You are transcribing a document image for a microbiology research archive. "
    "Extract ALL text verbatim. Then add a one-line description of any figure, "
    "chart, plate, or gel. Output plain text only."
)


def detect_kind(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "docx"
    if ext in IMAGE_EXTS:
        return "image"
    if ext in TEXT_EXTS:
        return "text"
    return "text"


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


async def _extract_image(path: Path) -> str:
    """Transcribe an image via the vision model (OCR + figure description)."""
    img_bytes = path.read_bytes()
    out: list[str] = []
    async for tok in ollama.chat_stream(
        model=settings.chat_model,
        messages=[{"role": "user", "content": _OCR_PROMPT}],
        images=[img_bytes],
    ):
        out.append(tok)
    return "".join(out)


async def extract_text(path: Path, kind: str) -> str:
    if kind == "pdf":
        return _extract_pdf(path)
    if kind == "docx":
        return _extract_docx(path)
    if kind == "image":
        return await _extract_image(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def chunk_text(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    size, overlap = settings.chunk_size, settings.chunk_overlap
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + size
        # prefer to break on a paragraph/sentence boundary near the window edge
        if end < len(text):
            boundary = text.rfind("\n", start + size // 2, end)
            if boundary == -1:
                boundary = text.rfind(". ", start + size // 2, end)
            if boundary != -1:
                end = boundary + 1
        chunks.append(text[start:end].strip())
        start = max(end - overlap, start + 1)
    return [c for c in chunks if c]


async def ingest_document(doc_id: int, path: Path, kind: str, topic: str) -> int:
    """Full pipeline for one document. Returns number of chunks indexed.

    Marks the document row ``indexed`` or ``error`` in SQLite.
    """
    try:
        text = await extract_text(path, kind)
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("No extractable text found in document.")
        embeddings = [await ollama.embed(c) for c in chunks]
        vectorstore.add_chunks(doc_id, chunks, embeddings, topic, path.name)
        with tx() as conn:
            conn.execute(
                "UPDATE documents SET status='indexed', chunks=? WHERE id=?",
                (len(chunks), doc_id),
            )
        return len(chunks)
    except Exception as exc:  # noqa: BLE001
        with tx() as conn:
            conn.execute(
                "UPDATE documents SET status='error', error=? WHERE id=?",
                (str(exc), doc_id),
            )
        raise
